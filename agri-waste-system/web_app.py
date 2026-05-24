from __future__ import annotations

import io
import json
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Any, Optional

import pandas as pd
from fastapi import FastAPI, HTTPException, Request, File, UploadFile, Form
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, Field
from scipy.optimize import linprog

from data_loader import load_excel_data
from question_bank import question_bank
from recommender import RouteRecommender
from carbon_market import router as carbon_market_router, init_carbon_market

ROOT = Path(__file__).resolve().parent
TEMPLATES_DIR = ROOT / "web" / "templates"
STATIC_DIR = ROOT / "web" / "static"

# 确保目录存在
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
STATIC_DIR.mkdir(parents=True, exist_ok=True)

# 唯一且正确的模板初始化，禁用缓存避免 key 类型错误
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))
templates.env.cache_size = 0


# 使用 lifespan 替代 on_event
@asynccontextmanager
async def lifespan(app: FastAPI):
    # --- 启动时执行 ---
    routes_df, rules_df = load_excel_data(str(ROOT / "农副产品加工路线+红线规则.xlsx"))
    if routes_df is None or rules_df is None:
        raise RuntimeError("Excel 数据加载失败，请检查文件是否存在且可读")
    app.state.routes_df = routes_df
    app.state.rules_df = rules_df
    app.state.recommender = RouteRecommender(routes_df, rules_df)
    print("✅ 推荐器初始化完成，多目标模型加载成功")
    # 初始化碳积分市场数据库
    init_carbon_market()
    print("✅ 碳积分市场初始化完成")
    yield
    # --- 关闭时执行（可选）---
    print("🛑 应用正在关闭，清理资源...")


app = FastAPI(title="副产物处置方案智能推荐系统", version="0.2.0", lifespan=lifespan)
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
app.include_router(carbon_market_router)


class RecommendRequest(BaseModel):
    user_profile: dict[str, Any] = Field(default_factory=dict)
    weights: Optional[dict[str, float]] = None


# 默认路线年处理能力上限（吨/年）
DEFAULT_CAPACITIES = {
    "茶渣制生物炭/吸附材料": 500,
    "茶渣发酵生产有机肥": 800,
    "茶渣提取茶多酚": 200,
    "果渣提取果胶": 300,
    "果渣生产饲料": 600,
    "果渣制备酵素": 150,
    "花生壳制备活性炭": 400,
    "花生壳生产板材": 500,
    "豆制品副产物生产蛋白饲料": 700,
    "豆制品副产物厌氧发酵产沼气": 600,
    "中药残余提取黄酮": 200,
    "中药渣制备生物质颗粒燃料": 800,
    "中药渣堆肥生产土壤改良剂": 500,
}


def get_route_capacity(route_name: str) -> float:
    """获取路线处理能力上限，默认300吨"""
    return DEFAULT_CAPACITIES.get(route_name, 300)


# ==================== 页面路由（手动渲染模板，绕过 TemplateResponse 缓存 bug） ====================
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    template = templates.get_template("index.html")
    content = template.render({"request": request, "title": "工艺知识库 + 数据驱动推荐"})
    return HTMLResponse(content=content)


@app.get("/grading", response_class=HTMLResponse)
def grading(request: Request):
    grading_dir = STATIC_DIR / "grading"
    grading_images = []
    if grading_dir.exists() and grading_dir.is_dir():
        for fp in sorted(grading_dir.iterdir()):
            if fp.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
                grading_images.append(
                    {
                        "filename": fp.name,
                        "url": f"/static/grading/{fp.name}",
                    }
                )
    template = templates.get_template("grading.html")
    content = template.render({
        "request": request,
        "title": "看图分级教学",
        "grading_images": grading_images,
    })
    return HTMLResponse(content=content)


@app.get("/recommend", response_class=HTMLResponse)
def recommend_page(request: Request):
    template = templates.get_template("recommend.html")
    content = template.render({"request": request, "title": "智能推荐"})
    return HTMLResponse(content=content)


@app.get("/batch-recommend", response_class=HTMLResponse)
def batch_recommend_page(request: Request):
    template = templates.get_template("batch_recommend.html")
    content = template.render({"request": request, "title": "批量推荐与供应链匹配"})
    return HTMLResponse(content=content)


@app.get("/carbon-market", response_class=HTMLResponse)
def carbon_market_page(request: Request):
    template = templates.get_template("carbon_market.html")
    content = template.render({"request": request, "title": "碳积分模拟交易平台"})
    return HTMLResponse(content=content)


# ==================== API 路由 ====================
@app.get("/api/options")
def get_options():
    return {"question_bank": question_bank}


@app.get("/api/routes")
def list_routes(material: str | None = None):
    df = app.state.routes_df
    if material:
        recommender: RouteRecommender = app.state.recommender
        df = df[df.apply(lambda x: recommender.is_material_match(material, x.get("适用原料类")), axis=1)].copy()
    cols = [c for c in ["路线编号", "路线名称", "适用原料类", "适用等级", "首选目标", "风险等级", "推荐强度", "推荐理由关键词"] if c in df.columns]
    return {"routes": df[cols].fillna("").to_dict(orient="records")}


@app.post("/api/recommend")
def recommend_api(payload: RecommendRequest):
    recommender: RouteRecommender = app.state.recommender
    profile = recommender.normalize_user_profile(payload.user_profile)

    candidate_routes = recommender.get_candidate_routes(profile)
    if candidate_routes.empty:
        raise HTTPException(status_code=400, detail="无可用合规路线，请检查原料情况")

    results = []
    for idx, route in candidate_routes.iterrows():
        scores = recommender.predict_all_scores(route, profile)
        item = route.to_dict()
        item.update(scores)
        results.append({"data": item, "route_obj": route})

    weights = payload.weights
    if weights is not None:
        total = sum(weights.values())
        if total > 0:
            weights = {k: v/total for k, v in weights.items()}
        for res in results:
            item = res["data"]
            weighted = (weights.get("成本", 0) * item.get("成本得分", 0) +
                        weights.get("技术", 0) * item.get("技术得分", 0) +
                        weights.get("碳减排", 0) * item.get("碳减排得分", 0) +
                        weights.get("市场", 0) * item.get("市场得分", 0))
            item["加权总分"] = weighted
        results.sort(key=lambda x: x["data"]["加权总分"], reverse=True)
    else:
        results.sort(key=lambda x: x["data"]["匹配得分"], reverse=True)

    top_results = results[:3]
    final_results = []
    for res in top_results:
        item = res["data"]
        route_obj = res["route_obj"]
        explanation = recommender.explain_route(route_obj, profile)
        final_results.append({**item, "explanation": explanation})

    return {"user_profile": profile, "recommendations": final_results, "weights": weights}


@app.post("/api/batch_preview")
async def batch_preview(file: UploadFile = File(...)):
    """预览上传文件，返回可能涉及的路线列表（用于产能配置）"""
    recommender: RouteRecommender = app.state.recommender
    contents = await file.read()
    try:
        if file.filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(contents))
        else:
            df = pd.read_excel(io.BytesIO(contents))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件解析失败: {str(e)}")

    all_routes = set()
    for _, row in df.iterrows():
        profile = {
            "原料类型": row.get("原料类型", ""),
            "霉变情况": row.get("霉变情况", "无"),
            "来源情况": row.get("来源情况", "明确可追溯"),
            "预估等级": row.get("预估等级", "良"),
        }
        candidate_routes = recommender.get_candidate_routes(profile)
        for _, route in candidate_routes.iterrows():
            all_routes.add(route.get("路线名称", ""))
    return {"routes": list(all_routes)}


@app.post("/api/batch_recommend")
async def batch_recommend(
    file: UploadFile = File(...),
    capacities_json: str = Form(None)
):
    """
    批量推荐与供应链匹配
    上传Excel/CSV文件，包含字段：
    - 原料类型, 霉变情况, 来源情况, 预估等级, 预计吨数
    - 专属字段：根据原料类型提供（茶渣_储存时间, 茶渣_是否含调味, 果渣_褐变情况等）
    """
    recommender: RouteRecommender = app.state.recommender

    # 1. 读取上传文件
    contents = await file.read()
    try:
        if file.filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(contents))
        else:
            df = pd.read_excel(io.BytesIO(contents))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件解析失败: {str(e)}")

    # 2. 检查必需列
    required_cols = ["原料类型", "霉变情况", "来源情况", "预估等级", "预计吨数"]
    missing = [c for c in required_cols if c not in df.columns]
    if missing:
        raise HTTPException(status_code=400, detail=f"缺少必需列: {missing}")

    # 3. 处理每个批次，获取候选路线及得分
    batches = []
    batch_candidates = []

    for idx, row in df.iterrows():
        profile = {
            "原料类型": row["原料类型"],
            "霉变情况": row["霉变情况"],
            "来源情况": row["来源情况"],
            "预估等级": row["预估等级"],
            "预计吨数": float(row["预计吨数"]),
        }
        material = row["原料类型"]
        if material == "茶渣":
            profile["储存时间"] = row.get("茶渣_储存时间", "否")
            profile["是否含调味"] = row.get("茶渣_是否含调味", "否")
        elif material == "果渣（柑橘/葡萄/猕猴桃等）":
            profile["褐变情况"] = row.get("果渣_褐变情况", "否")
        elif material == "花生壳/板栗壳":
            profile["泥沙情况"] = row.get("泥沙情况", "否")
        elif material == "豆制品/粮食发酵副产物":
            profile["含水率情况"] = row.get("含水率情况", "否")
        elif material == "中药残余废弃物":
            profile["药材明确度"] = row.get("药材明确度", "是")

        batches.append({
            "batch_id": idx,
            "profile": profile,
            "tonnage": profile["预计吨数"]
        })

        candidate_routes = recommender.get_candidate_routes(profile)
        if candidate_routes.empty:
            batch_candidates.append([])
            continue

        route_scores = []
        for _, route in candidate_routes.iterrows():
            scores = recommender.predict_all_scores(route, profile)
            match_score = scores.get("匹配得分", 0)
            route_scores.append({
                "route_id": route.get("路线编号", ""),
                "route_name": route.get("路线名称", ""),
                "match_score": match_score,
                "capacity": get_route_capacity(route.get("路线名称", "")),
            })
        route_scores.sort(key=lambda x: x["match_score"], reverse=True)
        batch_candidates.append(route_scores)

    # 4. 解析自定义产能
    custom_capacities = {}
    if capacities_json:
        try:
            custom_capacities = json.loads(capacities_json)
        except:
            pass

    # 5. 收集所有唯一的路线
    all_routes = {}
    for candidates in batch_candidates:
        for c in candidates:
            rname = c["route_name"]
            if rname not in all_routes:
                capacity = custom_capacities.get(rname, c["capacity"])
                all_routes[rname] = {
                    "capacity": capacity,
                    "route_id": c["route_id"]
                }

    route_names = list(all_routes.keys())
    n_batches = len(batches)
    n_routes = len(route_names)

    # 构建决策变量
    c_coeff = []
    variable_info = []

    for i, candidates in enumerate(batch_candidates):
        score_map = {c["route_name"]: c["match_score"] for c in candidates}
        for j, rname in enumerate(route_names):
            score = score_map.get(rname, 0)
            c_coeff.append(-score)  # linprog求最小，取负
            variable_info.append((i, rname, j))

    n_vars = len(c_coeff)

    # 约束：批次约束 + 路线约束
    constraints_eq = []
    b_ub = []

    # 批次约束 (每个批次分配量 <= 预计吨数)
    for i in range(n_batches):
        row = [0] * n_vars
        for var_idx, (batch_idx, rname, _) in enumerate(variable_info):
            if batch_idx == i:
                row[var_idx] = 1
        constraints_eq.append(row)
        b_ub.append(batches[i]["tonnage"])

    # 路线约束 (每条路线分配总量 <= 产能上限)
    for j, rname in enumerate(route_names):
        row = [0] * n_vars
        for var_idx, (_, _, route_j) in enumerate(variable_info):
            if route_j == j:
                row[var_idx] = 1
        constraints_eq.append(row)
        b_ub.append(all_routes[rname]["capacity"])

    bounds = [(0, None)] * n_vars

    # 求解线性规划
    try:
        result = linprog(c_coeff, A_ub=constraints_eq, b_ub=b_ub, bounds=bounds, method='highs')
        if not result.success:
            raise HTTPException(status_code=400, detail=f"线性规划无可行解: {result.message}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"优化求解失败: {str(e)}")

    # 6. 解析结果
    allocations = []
    total_score = 0
    route_usage = {rname: 0 for rname in route_names}

    for var_idx, (batch_idx, rname, _) in enumerate(variable_info):
        amount = result.x[var_idx]
        if amount > 0.01:
            score_per_ton = None
            for c in batch_candidates[batch_idx]:
                if c["route_name"] == rname:
                    score_per_ton = c["match_score"]
                    break
            if score_per_ton is not None:
                total_score += score_per_ton * amount
                route_usage[rname] += amount
                allocations.append({
                    "batch_id": batch_idx,
                    "原料类型": batches[batch_idx]["profile"]["原料类型"],
                    "预计吨数": batches[batch_idx]["tonnage"],
                    "分配吨数": round(amount, 2),
                    "路线名称": rname,
                    "匹配得分": score_per_ton,
                    "贡献效益": round(score_per_ton * amount, 2),
                })

    # 7. 构建路线产能汇总
    route_capacity_summary = []
    for rname in route_names:
        route_capacity_summary.append({
            "路线名称": rname,
            "总产能上限": all_routes[rname]["capacity"],
            "已分配量": round(route_usage[rname], 2),
            "剩余产能": round(all_routes[rname]["capacity"] - route_usage[rname], 2),
        })

    return {
        "success": True,
        "total_batches": n_batches,
        "total_allocated_tonnage": round(sum(a["分配吨数"] for a in allocations), 2),
        "total_benefit_score": round(total_score, 2),
        "allocations": allocations,
        "route_capacities": route_capacity_summary,
        "batch_details": [
            {
                "batch_id": b["batch_id"],
                "原料类型": b["profile"]["原料类型"],
                "预计吨数": b["tonnage"],
                "已分配吨数": round(sum(a["分配吨数"] for a in allocations if a["batch_id"] == b["batch_id"]), 2),
            }
            for b in batches
        ],
    }


# ==================== 碳足迹计算 API ====================
import json
from carbon_market import add_carbon_points

# 加载碳排放因子
CARBON_FACTORS_PATH = ROOT / "carbon_factors.json"
with open(CARBON_FACTORS_PATH, "r", encoding="utf-8") as f:
    CARBON_FACTORS = json.load(f)

class CarbonRequest(BaseModel):
    material: str
    tonnage: float
    transport_distance: float = 100
    transport_mode: str = "truck"
    packaging_type: str = "散装"
    preprocess_energy: str = "电"
    process_route: str = "茶渣制生物炭/吸附材料"
    user_id: str = "guest"  # 新增：用于碳积分归属

@app.post("/api/carbon_footprint")
def calculate_carbon_footprint(req: CarbonRequest):
    # 运输排放
    transport_factor = CARBON_FACTORS["transport"].get(req.transport_mode, 0.1)
    transport_emission = req.transport_distance * req.tonnage * transport_factor
    # 包装排放
    packaging_factor = CARBON_FACTORS["packaging"].get(req.packaging_type, 0)
    packaging_emission = req.tonnage * packaging_factor
    # 预处理排放
    preprocess_factor = CARBON_FACTORS["preprocess"].get(req.preprocess_energy, 0)
    preprocess_emission = req.tonnage * preprocess_factor
    # 工艺排放
    process_factor = CARBON_FACTORS["process"].get(req.process_route, CARBON_FACTORS["process"]["default"])
    process_emission = req.tonnage * process_factor

    total_emission = transport_emission + packaging_emission + preprocess_emission + process_emission

    # 基准线排放（填埋：假设每吨原料产生 500 kg CO₂e）
    baseline_emission = req.tonnage * 500
    reduction = max(0, baseline_emission - total_emission)  # 减排量

    # 发放碳积分（1 kg 减排 = 1 积分）
    if reduction > 0:
        add_carbon_points(req.user_id, reduction, "carbon_footprint", reduction)

    return {
        "total_emission": round(total_emission, 2),
        "transport_emission": round(transport_emission, 2),
        "packaging_emission": round(packaging_emission, 2),
        "preprocess_emission": round(preprocess_emission, 2),
        "process_emission": round(process_emission, 2),
        "baseline_emission": round(baseline_emission, 2),
        "reduction": round(reduction, 2),
        "carbon_points_awarded": round(reduction, 2) if reduction > 0 else 0,
        "unit": "kg CO₂e",
        "details": {
            "transport_factor": transport_factor,
            "packaging_factor": packaging_factor,
            "preprocess_factor": preprocess_factor,
            "process_factor": process_factor
        }
    }


@app.get("/carbon-calculator", response_class=HTMLResponse)
def carbon_calculator_page(request: Request):
    template = templates.get_template("carbon_calculator.html")
    content = template.render({"request": request, "title": "碳足迹计算器"})
    return HTMLResponse(content=content)


# ==================== 副产物资源价值评估工具 ====================
VALUATION_RULES_PATH = ROOT / "valuation_rules.json"
with open(VALUATION_RULES_PATH, "r", encoding="utf-8") as f:
    valuation_rules = json.load(f)

class ValuationRequest(BaseModel):
    material: str
    origin: str
    season: str
    moisture: str
    impurity: str

@app.post("/api/valuate")
def evaluate_resource(req: ValuationRequest):
    base = valuation_rules["base_scores"].get(req.material, valuation_rules["base_scores"]["茶渣"])
    mo = valuation_rules["moisture_adjust"].get(req.moisture, {})
    imp = valuation_rules["impurity_adjust"].get(req.impurity, {})
    sea = valuation_rules["season_adjust"].get(req.season, {})
    org = valuation_rules["origin_adjust"].get(req.origin, {})
    
    scores = {}
    for key in ["economic", "environmental", "energy", "feasibility"]:
        score = base[key] + mo.get(key, 0) + imp.get(key, 0) + sea.get(key, 0) + org.get(key, 0)
        scores[key] = max(0, min(10, score))
    
    # 简化建议逻辑
    if scores["economic"] >= 7 and scores["feasibility"] >= 7:
        suggestion = valuation_rules["disposal_suggestions"]["经济价值高且可行性高"]
    elif scores["environmental"] >= 7:
        suggestion = valuation_rules["disposal_suggestions"]["环保价值高"]
    elif scores["energy"] >= 7:
        suggestion = valuation_rules["disposal_suggestions"]["能源价值高"]
    elif scores["feasibility"] >= 7 and scores["economic"] < 6:
        suggestion = valuation_rules["disposal_suggestions"]["可行性高但经济一般"]
    else:
        suggestion = valuation_rules["disposal_suggestions"]["综合偏低"]
    
    if req.moisture == "高":
        suggestion += " 注意：原料含水量高，需增加脱水预处理成本。"
    if req.impurity == "高":
        suggestion += " 注意：杂质含量高，需增加筛分/净化步骤。"
    
    return {
        "scores": scores,
        "suggestion": suggestion,
        "radar_data": [scores["economic"], scores["environmental"], scores["energy"], scores["feasibility"]],
        "labels": ["经济价值", "环保价值", "能源价值", "技术可行性"]
    }

@app.get("/resource-valuation", response_class=HTMLResponse)
def valuation_page(request: Request):
    template = templates.get_template("valuation.html")
    content = template.render({"request": request, "title": "副产物资源价值评估"})
    return HTMLResponse(content=content)


# ==================== 合同生成 API ====================
from fastapi.responses import Response
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

class ContractRequest(BaseModel):
    material: str
    route_name: str
    tonnage: float
    price_per_ton: float = 0
    quality_requirements: str = ""
    transport_mode: str = "货车"
    delivery_location: str = "需方指定地点"
    payment_terms: str = "货到付款"
    contract_period: str = "合同签订后30天内"

@app.post("/api/generate_contract")
def generate_contract(req: ContractRequest):
    doc = Document()
    title = doc.add_heading('农副产物高值化处置合同', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contract_no = f"BVC-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
    doc.add_paragraph(f"合同编号：{contract_no}")
    doc.add_paragraph(f"签订日期：{datetime.datetime.now().strftime('%Y年%m月%d日')}")
    doc.add_paragraph()
    doc.add_paragraph("甲方（供应方）：____________________")
    doc.add_paragraph("乙方（接收方）：____________________")
    doc.add_paragraph()
    doc.add_heading('第一条 原料信息', level=1)
    doc.add_paragraph(f"1.1 原料类型：{req.material}")
    doc.add_paragraph(f"1.2 数量：{req.tonnage} 吨")
    doc.add_paragraph(f"1.3 单价：{req.price_per_ton} 元/吨（含税）")
    doc.add_paragraph(f"1.4 总金额：{req.price_per_ton * req.tonnage} 元")
    doc.add_heading('第二条 质量要求', level=1)
    if req.quality_requirements:
        doc.add_paragraph(req.quality_requirements)
    else:
        doc.add_paragraph("2.1 原料应符合国家相关标准及行业规范。")
        doc.add_paragraph("2.2 霉变率不得超过5%，杂质率不得超过3%。")
        doc.add_paragraph("2.3 具体技术指标参照《农副产物处置质量验收细则》。")
    doc.add_heading('第三条 处置路线与工艺', level=1)
    doc.add_paragraph(f"3.1 乙方承诺按照《{req.route_name}》工艺路线进行处置。")
    doc.add_paragraph("3.2 处置过程应符合环保要求，碳排放需控制在约定范围内。")
    doc.add_heading('第四条 运输与交付', level=1)
    doc.add_paragraph(f"4.1 运输方式：{req.transport_mode}")
    doc.add_paragraph(f"4.2 交付地点：{req.delivery_location}")
    doc.add_paragraph("4.3 运输费用由甲方承担，乙方协助卸货。")
    doc.add_heading('第五条 付款方式', level=1)
    doc.add_paragraph(f"5.1 {req.payment_terms}")
    doc.add_paragraph("5.2 乙方收到发票后15个工作日内支付全款。")
    doc.add_heading('第六条 违约责任', level=1)
    doc.add_paragraph("6.1 甲方所供原料不符合质量要求的，乙方有权拒收或要求降价处理。")
    doc.add_paragraph("6.2 乙方未按约定付款的，每逾期一日，按未付金额的0.05%支付违约金。")
    doc.add_heading('第七条 争议解决', level=1)
    doc.add_paragraph("7.1 本合同履行中发生争议，双方应协商解决；协商不成的，提交乙方所在地人民法院诉讼解决。")
    doc.add_heading('第八条 其他约定', level=1)
    doc.add_paragraph(f"8.1 合同有效期：{req.contract_period}")
    doc.add_paragraph("8.2 本合同一式两份，甲乙双方各执一份，具有同等法律效力。")
    doc.add_paragraph("8.3 本合同自双方签字盖章之日起生效。")
    doc.add_paragraph()
    doc.add_paragraph("甲方（盖章）：________________    乙方（盖章）：________________")
    doc.add_paragraph("代表签字：________________        代表签字：________________")
    doc.add_paragraph("日期：________________            日期：________________")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=contract_{contract_no}.docx"}
    )

@app.get("/contract-generator", response_class=HTMLResponse)
def contract_generator_page(request: Request):
    template = templates.get_template("contract_generator.html")
    content = template.render({"request": request, "title": "智能合同生成器"})
    return HTMLResponse(content=content)